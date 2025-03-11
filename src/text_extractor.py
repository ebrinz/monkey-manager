import os
import sys
import json
import fitz  # PyMuPDF
import datetime
from oletools.olevba import VBA_Parser
from docx import Document
from forensic_logger import ForensicLogger
from file_renamer import FilenamingUtility

class TextExtractor:
    def __init__(self, input_dir, output_dir, mapping_file="/mapping/mapping_file.xlsx", force_reprocess=False):
        self.input_dir = input_dir
        self.output_dir = output_dir
        self.logger = ForensicLogger("/output/logs")
        self.force_reprocess = force_reprocess
        
        # Initialize filename utility with mapping file
        enable_renaming = os.environ.get('ENABLE_FILE_RENAMING', 'true').lower() == 'true'
        self.filename_util = FilenamingUtility(mapping_file, enable_renaming)
        
        os.makedirs(output_dir, exist_ok=True)

    def process_files(self):
        """Process all files in the input directory."""
        self.logger.log_system_state()

        for root, _, files in os.walk(self.input_dir):
            for fname in files:
                if fname.startswith('.') or fname in [".DS_Store", ".gitkeep"]:
                    continue

                file_path = os.path.join(root, fname)
                base, ext = os.path.splitext(fname)
                ext_lower = ext.lower()
                
                # Skip unsupported file types
                if ext_lower not in ['.pdf', '.docx']:
                    continue
                
                # Get output filename and respondent info
                new_filename, respondent_id, col_num = self.filename_util.get_output_filename(fname, '.json')
                json_path = os.path.join(self.output_dir, str(new_filename))
                
                # Skip if output already exists (unless force flag is set)
                if os.path.exists(json_path) and not self.force_reprocess:
                    self.logger.log_file_event('skip_existing_output', file_path, {
                        'output_path': json_path
                    })
                    continue

                # Process based on extension
                text = None
                if ext_lower == '.pdf':
                    text = self.sanitize_pdf(file_path)
                elif ext_lower == '.docx':
                    text = self.sanitize_docx(file_path)

                # Save extracted text as JSON
                if text:
                    # Create JSON document
                    doc = {
                        'text': text,
                        'filename': fname,
                        'filetype': ext_lower.lstrip('.'),
                        'original_path': file_path,
                        'extraction_timestamp': datetime.datetime.now().isoformat()
                    }
                    
                    # Add respondent info if available from mapping
                    if respondent_id:
                        doc['respondent_id'] = respondent_id
                        doc['file_column'] = col_num
                    
                    # Save as JSON
                    with open(json_path, "w", encoding="utf-8") as out_f:
                        json.dump(doc, out_f, ensure_ascii=False, indent=2)
                    
                    self.logger.log_file_event('json_saved', json_path, {
                        'original_file': file_path,
                        'text_length': len(text)
                    })

    def sanitize_pdf(self, input_file):
        """Extract text from PDF while logging operations."""
        try:
            self.logger.log_file_event('pdf_start', input_file)
            doc = fitz.open(input_file)
            text = ""
            for page in doc:
                text += page.get_text()
            self.logger.log_file_event('pdf_success', input_file, {
                'pages': len(doc),
                'text_length': len(text)
            })
            return text
        except Exception as e:
            self.logger.log_anomaly('pdf_error', {
                'file': input_file,
                'error': str(e)
            })
            return None

    def sanitize_docx(self, input_file):
        """Extract text from DOCX while checking for macros."""
        try:
            self.logger.log_file_event('docx_start', input_file)
            
            # Check for macros
            vba_parser = VBA_Parser(input_file)
            has_macros = vba_parser.detect_vba_macros()
            if has_macros:
                self.logger.log_anomaly('macro_detected', {
                    'file': input_file
                })
            
            # Extract text
            doc = Document(input_file)
            text = "\n".join(p.text for p in doc.paragraphs)
            
            self.logger.log_file_event('docx_success', input_file, {
                'paragraphs': len(doc.paragraphs),
                'has_macros': has_macros,
                'text_length': len(text)
            })
            return text
        except Exception as e:
            self.logger.log_anomaly('docx_error', {
                'file': input_file,
                'error': str(e)
            })
            return None


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Extract text from PDF and DOCX files.")
    parser.add_argument("input_dir", help="Directory containing files to process")
    parser.add_argument("output_dir", help="Directory for JSON outputs")
    parser.add_argument("--force", "-f", action="store_true", help="Force reprocessing of files that already have outputs")
    parser.add_argument("--mapping", "-m", help="Path to mapping file for filename conversion")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input_dir):
        print(f"Error: Input directory '{args.input_dir}' does not exist.")
        sys.exit(1)

    extractor = TextExtractor(
        args.input_dir, 
        args.output_dir, 
        mapping_file=args.mapping if args.mapping else "/mapping/mapping_file.xlsx",
        force_reprocess=args.force
    )
    extractor.process_files()